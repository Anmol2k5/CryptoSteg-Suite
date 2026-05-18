import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def create_report():
    doc = docx.Document()

    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Title Page
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_inst = p_inst.add_run("Institute of Engineering and Technology (IET)")
    run_inst.font.size = Pt(20)
    run_inst.font.bold = True
    run_inst.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Cryptography Project Report\n")
    run_sub.font.size = Pt(16)
    run_sub.font.bold = True

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("Project Title: Secure Steganographic AI Toolkit (SSAT) & CryptoSteg Suite\n\n\n")
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x00, 0x80, 0x80)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run(
        "PREPARED BY:\nXYZ (2023BTech000)\nABC (2023BTech000)\n\n\n"
        "FACULTY GUIDE:\nDr. Assistant Professor\nDepartment of Computer Science & Engineering\n\n\n"
        "March 2026"
    )
    run_meta.font.size = Pt(12)
    run_meta.font.bold = True

    doc.add_page_break()

    # Table of Contents
    h_toc = doc.add_heading("TABLE OF CONTENTS", level=1)
    h_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc_items = [
        "1. Abstract",
        "2. Introduction",
        "3. Problem Statement",
        "4. Methodology and Theoretical Background",
        "5. Architecture and Key Design Decisions",
        "   5.1 High-Level System Design",
        "   5.2 Tech Stack",
        "   5.3 Additional Design & Serverless Bridge",
        "6. Work Completed",
        "   6.1 Research and Algorithm Design",
        "   6.2 Backend & Cryptographic Vault",
        "   6.3 Frontend & Visual Analytics Core",
        "   6.4 Autonomous AI Agent & MCP Integration",
        "7. Plan for Remaining Project Timeline",
        "8. Risks and Mitigation Strategies",
        "9. Appendix",
        "   A.1 Key Code Snippets",
        "   A.2 CLI Usage Guide",
        "10. References"
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # Section 1: Abstract
    doc.add_heading("1. Abstract", level=1)
    doc.add_paragraph(
        "The rapid adoption of multimodal Vision-Language Models (VLMs) and autonomous AI agents has introduced a novel cyber-attack vector: Indirect Prompt Injection (IPI) via digital media steganography. Malicious actors can embed covert instructions inside the Least Significant Bits (LSBs) of images, hijacking AI behavior upon ingestion. This report presents the Secure Steganographic AI Toolkit (SSAT), a robust dual-purpose forensic defense and digital provenance framework. SSAT integrates real-time bitwise decontamination filters, advanced statistical steganalysis (Chi-Squared PoV, Sample Pair Analysis, RS Flipper test, Shannon Entropy mapping), frequency-domain Discrete Wavelet Transform (DWT) watermarking, military-grade AES-256 encryption, and full Model Context Protocol (MCP) server support. Coupled with an enterprise-grade web dashboard and ReportLab forensic PDF generation, SSAT provides zero-trust multimodal verification for modern enterprise and autonomous AI architectures."
    )

    # Section 2: Introduction
    doc.add_heading("2. Introduction", level=1)
    doc.add_paragraph(
        "In recent years, artificial intelligence has transitioned from text-only Large Language Models (LLMs) to multimodal systems capable of visual reasoning. While powerful, these systems process image pixel data as unconstrained input streams. Steganography—the practice of concealing secret data within ordinary non-secret digital media—historically served secure communication but now poses an active security threat to AI pipelines. An unsanitized image uploaded to an AI assistant or crawled by an autonomous web bot can contain imperceptible adversarial instructions.\n\n"
        "Furthermore, as AI agents autonomously generate valuable digital media (e.g., financial charts, architectural blueprints, synthetic datasets), verifying the authenticity and origin of these artifacts is paramount. The CryptoSteg Suite & SSAT project addresses both challenges by building an end-to-end cryptographic security suite that acts as an active firewall for incoming media and a non-repudiation signing authority for outgoing media."
    )

    # Section 3: Problem Statement
    doc.add_heading("3. Problem Statement", level=1)
    doc.add_paragraph(
        "1. Vulnerability of VLMs to Indirect Prompt Injections (IPI): Vision models process all pixel layers. When an image contains text payloads hidden via spatial LSB substitution, the model's visual reasoning layer extracts and executes the covert prompt override, circumventing text-based safety guardrails.\n"
        "2. Lack of Forensic Visibility in Media Streams: Traditional network firewalls cannot inspect image bit-planes. Security Operations Centers (SOCs) and AI engineers lack real-time statistical tools to detect covert steganographic C2 channels or data exfiltration.\n"
        "3. Absence of Provenance in Generative AI Output: High-value digital assets generated by autonomous bots can be copied, altered, or deepfaked without accountability. There is a lack of integrated spread-spectrum watermarking and traitor tracing to verify asset ownership."
    )

    # Section 4: Methodology and Theoretical Background
    doc.add_heading("4. Methodology and Theoretical Background", level=1)
    doc.add_heading("4.1 Least Significant Bit (LSB) Steganography", level=2)
    doc.add_paragraph(
        "Spatial LSB substitution replaces the k-th lowest bits of pixel color channels (RGB) with binary payload streams. In natural images, lower bit layers represent random thermal noise. Replacing them with encrypted ASCII or binary data alters the localized statistical distribution."
    )

    doc.add_heading("4.2 Chi-Squared Pairs of Values (PoV) Analysis", level=2)
    doc.add_paragraph(
        "In natural color distributions, pixel intensity occurrences 2i and 2i+1 vary organically. Spatial steganography forces these pairs into artificial statistical equilibrium. Computing the Chi-Squared variance against degrees of freedom yields a definitive tampering metric. Scores significantly exceeding 100 indicate high probability of LSB manipulation."
    )

    doc.add_heading("4.3 Sample Pair Analysis (SPA)", level=2)
    doc.add_paragraph(
        "Developed by Dumitrescu et al., SPA measures spatial correlation between adjacent pixel pairs across horizontal, vertical, and diagonal vectors. Natural images maintain smooth transitions (~50% correlation). Injecting random steganographic payloads disrupts this spatial correlation, allowing precise estimation of the embedded message length."
    )

    doc.add_heading("4.4 RS (Regular-Singular) Flipper Test", level=2)
    doc.add_paragraph(
        "Fridrich et al.'s RS analysis classifies pixel groups into Regular, Singular, and Unusable sets using positive and negative flipping masks. Unmodified images exhibit symmetry between these groups. Steganographic embedding breaks this symmetry. An RS ratio > 0.05 definitively proves spatial LSB inversion."
    )

    doc.add_heading("4.5 Discrete Wavelet Transform (DWT) & Cryptography", level=2)
    doc.add_paragraph(
        "To ensure watermark durability against resizing and JPEG compression, SSAT applies a 2D Haar wavelet decomposition, isolating media into four sub-bands (LL, LH, HL, HH). Payloads are encrypted using AES-256-GCM with PBKDF2HMAC key derivation (100,000 iterations) and embedded into the robust LH/HL coefficients."
    )

    # Section 5: Architecture and Key Design Decisions
    doc.add_heading("5. Architecture and Key Design Decisions", level=1)
    doc.add_heading("5.1 High Level System Design", level=2)
    doc.add_paragraph(
        "The system architecture is split into two primary pipelines: Inbound Defense (Zero-Trust VLM Firewall) and Outbound Provenance (Generative Watermarking Vault). Inbound images are analyzed via FastMCP; if infected, Level 2 bitwise scrubbing (& 0xFC) is applied before VLM ingestion. Outbound media generated by AI agents undergoes DWT decomposition and spread-spectrum signing before external delivery."
    )

    doc.add_heading("5.2 Tech Stack", level=2)
    doc.add_paragraph(
        "• Core Processing: Python 3.11+, OpenCV, Pillow, NumPy, PyWavelets\n"
        "• Cryptography: AES-256-GCM, PBKDF2HMAC, SHA-256\n"
        "• Web & Reporting: Flask, Tailwind CSS, Matplotlib, ReportLab\n"
        "• Agent Integration: Model Context Protocol (FastMCP)"
    )

    doc.add_heading("5.3 Additional Design", level=2)
    doc.add_paragraph(
        "The system incorporates a serverless cloud bridge (vercel.json and api/index.py) that exports the Flask application object. This decouples the processing engine from permanent infrastructure, allowing cloud providers like Vercel to dynamically instantiate Python workers to execute steganalysis and report generation on demand."
    )

    # Section 6: Work Completed
    doc.add_heading("6. Work Completed", level=1)
    doc.add_heading("6.1 Research and Design", level=2)
    doc.add_paragraph(
        "• Integrated standalone LSB encoder/decoder CLI with API support for Pollinations.ai carrier generation (model=flux).\n"
        "• Developed robust statistical modules (ssat/analyze.py) computing Chi-Squared metrics, SPA spatial correlation matrices, RS flipping symmetry ratios, and localized Shannon entropy blocks."
    )

    doc.add_heading("6.2 Backend", level=2)
    doc.add_paragraph(
        "• Engineered ssat/crypto.py establishing an impenetrable vault utilizing AES-256 authenticated encryption with GCM tag verification to prevent payload tampering.\n"
        "• Built zero-latency image decontamination filters (ssat/sanitize.py) capable of bitwise scrubbing across RGB channels."
    )

    doc.add_heading("6.3 Frontend", level=2)
    doc.add_paragraph(
        "• Developed a fully responsive web dashboard (ssat/visualize.py) featuring dark mode aesthetics, drag-and-drop uploads, real-time batch progress tracking, and interactive bit-plane deconstruction.\n"
        "• Integrated a professional PDF reporting engine (ReportLab) generating court-admissible dossiers complete with SHA-256 chain-of-custody verification."
    )

    doc.add_heading("6.4 Autonomous AI Agent & MCP Integration", level=2)
    doc.add_paragraph(
        "• Deployed an active Model Context Protocol server (ssat_mcp_server.py) exposing inspect_and_sanitize_image, sign_bot_output, and trace_leaked_bot_artifact tools for ClaudeBot and Modal cloud workers."
    )

    # Section 7: Plan For Remaining Project Timeline
    doc.add_heading("7. Plan For Remaining Project Timeline", level=1)
    doc.add_paragraph(
        "• Vectorized Tensor Processing (PyTorch/CuPy): Migration of spatial loop logic to GPU tensors for 50x speedup by late March.\n"
        "• Blind QIM DWT Extraction Engine: Implementation of Quantization Index Modulation for carrier-less watermark extraction by early April.\n"
        "• Asynchronous Celery/Redis Queue: High-throughput background processing for massive 50+ image forensic batches by mid-April."
    )

    # Section 8: Risks and Mitigation Strategies
    doc.add_heading("8. Risks and Mitigation Strategies", level=1)
    
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    headers = ["Risk / Threat Vector", "Potential Impact", "Implemented Mitigation Strategy"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "1B365D")
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    data = [
        ("High-Frequency Natural Textures", "False positive stego warnings on highly textured natural images (foliage, sand).", "Multi-layered consensus model: Flags require both Chi2 anomaly and RS/SPA spatial divergence before triggering alerts."),
        ("Adversarial Cropping / Scaling", "Destroys spatial LSB watermarks, rendering tracking impossible.", "DWT frequency embedding targets LH/HL sub-bands with redundant spread-spectrum tiling across image quadrants."),
        ("HTTP Timeout on Massive Batches", "Dashboard crashes when users upload 50+ high-resolution images.", "Image downsampling before preview rendering; future roadmap transitions batch processing to asynchronous task queues.")
    ]
    for row_idx, row_data in enumerate(data, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text

    # Section 9: Appendix
    doc.add_heading("9. Appendix", level=1)
    doc.add_heading("A.1 Core Code Snippets", level=2)
    doc.add_paragraph(
        "Sanitization Firewall Function:\n"
        "def sanitize_image(input_path, output_path, level=1):\n"
        "    img = cv2.imread(input_path, cv2.IMREAD_COLOR)\n"
        "    mask = 0xFE if level == 1 else 0xFC\n"
        "    sanitized = cv2.bitwise_and(img, mask)\n"
        "    cv2.imwrite(output_path, sanitized)\n"
        "    return output_path"
    )

    doc.add_heading("A.2 CLI Usage Guide", level=2)
    doc.add_paragraph(
        "• Gateway Decontamination: python main.py sanitize --input untrusted.png --output safe.png --level 2\n"
        "• Forensic Steganalysis: python main.py analyze --input payload.png --heatmap map.png --all\n"
        "• Web Dashboard: python main.py dashboard --host 127.0.0.1 --port 5000"
    )

    # Section 10: References
    doc.add_heading("10. References", level=1)
    doc.add_paragraph(
        "1. Fridrich, J., Goljan, M., & Du, R. (2001). Detecting LSB steganography in color and gray-scale images. IEEE Multimedia, 8(4), 22-28.\n"
        "2. Dumitrescu, S., Wu, X., & Wang, Z. (2003). Detection of LSB steganography via sample pair analysis. IEEE Transactions on Signal Processing, 51(7), 1995-2007.\n"
        "3. Westfeld, A., & Pfitzmann, A. (1999). Attacks on Steganographic Systems. Information Hiding (LNCS 1768), 61-76.\n"
        "4. Model Context Protocol Specification. (2025). FastMCP Server Architecture for Autonomous AI Agents. Standard Protocols."
    )

    output_path = "Cryptography_Project_Report.docx"
    doc.save(output_path)
    print(f"Successfully generated {output_path}")

if __name__ == "__main__":
    create_report()
